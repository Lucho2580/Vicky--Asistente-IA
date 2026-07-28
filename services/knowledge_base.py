import json
import math
import re
from pathlib import Path
from typing import Callable, Dict, List, Optional

from core.paths import TRAINING_DIR
from database.knowledge_store import KnowledgeStore
from models.training_file import TrainingFile

SUPPORTED_TEXT_EXTENSIONS = {".txt", ".md", ".csv", ".json", ".log", ".pdf", ".docx", ".xlsx", ".xls"}

MAX_CONTENT_LENGTH = 200_000
MIN_KEYWORD_LENGTH = 3
SEMANTIC_MATCH_THRESHOLD = 0.25  # similitud coseno mínima para considerar un documento relevante

_STOPWORDS = {
    "que", "cual", "cuales", "cuál", "cuáles", "es", "la", "el", "los", "las",
    "de", "del", "en", "y", "o", "un", "una", "unos", "unas", "para", "por",
    "con", "sin", "se", "su", "sus", "al", "lo", "le", "les", "mi", "mis",
    "tu", "tus", "como", "cómo", "cuando", "cuándo", "donde", "dónde", "qué",
    "the", "is", "are", "of", "to", "for", "in", "on", "and", "or", "a", "an",
}


class UnsupportedFileTypeError(Exception):
    pass


class DocumentExtractionError(Exception):
    """El archivo tiene una extensión soportada, pero no se pudo leer su contenido (PDF corrupto, escaneado sin texto, .docx dañado, etc.)."""


def extract_keywords(text: str) -> List[str]:
    words = re.findall(r"[\wáéíóúñü]+", text.lower(), flags=re.UNICODE)
    return [w for w in words if len(w) >= MIN_KEYWORD_LENGTH and w not in _STOPWORDS]


def friendly_name(filename: str) -> str:
    stem = filename.rsplit(".", 1)[0] if "." in filename else filename
    words = stem.replace("_", " ").replace("-", " ").split()
    return " ".join(w.capitalize() for w in words) if words else stem


def cosine_similarity(a: List[float], b: List[float]) -> float:
    if not a or not b or len(a) != len(b):
        return 0.0
    dot = sum(x * y for x, y in zip(a, b))
    norm_a = math.sqrt(sum(x * x for x in a))
    norm_b = math.sqrt(sum(y * y for y in b))
    if norm_a == 0.0 or norm_b == 0.0:
        return 0.0
    return dot / (norm_a * norm_b)


def extract_text_from_file(path: Path, extension: str) -> str:
    if extension == ".pdf":
        return _extract_text_from_pdf(path)
    if extension == ".docx":
        return _extract_text_from_docx(path)
    if extension in (".xlsx", ".xls"):
        return _extract_text_from_excel(path)
    return path.read_text(encoding="utf-8", errors="replace")


def _extract_text_from_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise DocumentExtractionError(
            "Falta el paquete 'pypdf' para leer archivos PDF (pip install pypdf)."
        ) from exc

    try:
        reader = PdfReader(str(path))
    except Exception as exc:
        raise DocumentExtractionError(f"No se pudo abrir el PDF: {exc}") from exc

    pages_text = []
    for page in reader.pages:
        try:
            pages_text.append(page.extract_text() or "")
        except Exception:
            continue

    text = "\n".join(pages_text).strip()
    if not text:
        raise DocumentExtractionError(
            "El PDF no tiene texto extraíble (probablemente es un escaneo/imagen sin OCR)."
        )
    return text


def _extract_text_from_docx(path: Path) -> str:
    try:
        import docx
    except ImportError as exc:
        raise DocumentExtractionError(
            "Falta el paquete 'python-docx' para leer archivos Word (pip install python-docx)."
        ) from exc

    try:
        document = docx.Document(str(path))
    except Exception as exc:
        raise DocumentExtractionError(f"No se pudo abrir el documento Word: {exc}") from exc

    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    text = "\n".join(parts).strip()
    if not text:
        raise DocumentExtractionError("El documento Word no tiene texto (¿está vacío?).")
    return text


_EXCEL_PREVIEW_MAX_ROWS = 50


def _extract_text_from_excel(path: Path) -> str:
    try:
        import pandas as pd
    except ImportError as exc:
        raise DocumentExtractionError(
            "Falta el paquete 'pandas' para leer archivos Excel (pip install pandas openpyxl)."
        ) from exc

    try:
        sheets = pd.read_excel(path, sheet_name=None)
    except Exception as exc:
        raise DocumentExtractionError(f"No se pudo abrir el archivo Excel: {exc}") from exc

    if not sheets:
        raise DocumentExtractionError("El archivo Excel no tiene hojas con datos.")

    sections = []
    for sheet_name, df in sheets.items():
        if df.empty:
            sections.append(f"Hoja «{sheet_name}»: sin datos.")
            continue
        sections.append(_describe_excel_sheet(sheet_name, df))

    text = "\n\n".join(sections).strip()
    if not text:
        raise DocumentExtractionError("El archivo Excel no tiene datos legibles.")
    return text


def _describe_excel_sheet(sheet_name: str, df) -> str:
    import pandas as pd

    n_rows, n_cols = df.shape
    lines = [
        f"Hoja «{sheet_name}»: {n_rows} filas x {n_cols} columnas.",
        f"Columnas: {', '.join(str(c) for c in df.columns)}",
    ]

    numeric_cols = list(df.select_dtypes(include="number").columns)
    if numeric_cols:
        lines.append("")
        lines.append("Estadísticas reales calculadas sobre TODAS las filas (no son una estimación):")
        for col in numeric_cols:
            series = df[col].dropna()
            if series.empty:
                continue
            lines.append(
                f"- {col}: suma={series.sum():,.2f} | promedio={series.mean():,.2f} | "
                f"mínimo={series.min():,.2f} | máximo={series.max():,.2f} | cantidad={series.count()}"
            )

    date_cols = [c for c in df.columns if pd.api.types.is_datetime64_any_dtype(df[c])]
    if not date_cols:
        for c in df.columns:
            if df[c].dtype == object:
                try:
                    parsed = pd.to_datetime(df[c], errors="coerce")
                except Exception:
                    continue
                if parsed.notna().mean() > 0.8:
                    date_cols.append(c)

    if date_cols and numeric_cols:
        date_col = date_cols[0]
        try:
            temp = df.copy()
            temp[date_col] = pd.to_datetime(temp[date_col], errors="coerce")
            temp = temp.dropna(subset=[date_col])
            if not temp.empty:
                temp["_periodo"] = temp[date_col].dt.to_period("M").astype(str)
                grouped = temp.groupby("_periodo")[numeric_cols].sum(numeric_only=True)
                lines.append("")
                lines.append(f"Totales agrupados por mes (columna de fecha: {date_col}):")
                for periodo, row in grouped.iterrows():
                    valores = ", ".join(f"{col}={row[col]:,.2f}" for col in numeric_cols)
                    lines.append(f"- {periodo}: {valores}")
        except Exception:
            pass

    preview_df = df.head(_EXCEL_PREVIEW_MAX_ROWS)
    lines.append("")
    if n_rows > _EXCEL_PREVIEW_MAX_ROWS:
        lines.append(
            f"Vista previa (primeras {_EXCEL_PREVIEW_MAX_ROWS} de {n_rows} filas totales; "
            f"las estadísticas de arriba SÍ corresponden a las {n_rows} filas completas):"
        )
    else:
        lines.append("Datos completos:")
    lines.append(preview_df.to_markdown(index=False))

    return "\n".join(lines)


class KnowledgeBase:

    def __init__(self, store: Optional[KnowledgeStore] = None) -> None:
        self._store = store or KnowledgeStore()

    def add_document(self, file_path: str) -> TrainingFile:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"No se encontró el archivo: {file_path}")

        extension = path.suffix.lower()
        if extension not in SUPPORTED_TEXT_EXTENSIONS:
            supported = ", ".join(sorted(SUPPORTED_TEXT_EXTENSIONS))
            raise UnsupportedFileTypeError(
                f"Tipo de archivo '{extension or 'sin extensión'}' no soportado todavía. "
                f"Por ahora se aceptan: {supported}"
            )

        content = extract_text_from_file(path, extension)[:MAX_CONTENT_LENGTH]
        size_bytes = path.stat().st_size

        return self._store.add_training_file(
            filename=path.name,
            file_type=extension.lstrip("."),
            size_bytes=size_bytes,
            content_text=content,
        )

    def read_ephemeral_attachment(self, file_path: str) -> "tuple[str, str]":
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"No se encontró el archivo: {file_path}")

        extension = path.suffix.lower()
        if extension not in SUPPORTED_TEXT_EXTENSIONS:
            supported = ", ".join(sorted(SUPPORTED_TEXT_EXTENSIONS))
            raise UnsupportedFileTypeError(
                f"Tipo de archivo '{extension or 'sin extensión'}' no soportado todavía. "
                f"Por ahora se aceptan: {supported}"
            )

        content = extract_text_from_file(path, extension)[:MAX_CONTENT_LENGTH]
        return path.name, content

    def list_documents(self) -> List[TrainingFile]:
        return self._store.list_training_files()

    def sync_training_folder(self, folder_path: Optional[Path] = None) -> Dict[str, object]:
        folder = Path(folder_path) if folder_path else TRAINING_DIR
        folder.mkdir(parents=True, exist_ok=True)

        summary: Dict[str, object] = {"added": 0, "updated": 0, "removed": 0, "errors": []}

        disk_files = {
            str(p.resolve()): p
            for p in folder.iterdir()
            if p.is_file() and p.suffix.lower() in SUPPORTED_TEXT_EXTENSIONS
        }

        tracked_by_path = {doc.source_path: doc for doc in self._store.list_training_files_from_folder()}

        for path_str, path_obj in disk_files.items():
            try:
                mtime = path_obj.stat().st_mtime
                existing = tracked_by_path.get(path_str)
                if existing is None:
                    self._add_from_training_folder(path_obj, mtime)
                    summary["added"] += 1
                elif mtime > existing.source_mtime:
                    self._store.remove_training_file(existing.id)
                    self._add_from_training_folder(path_obj, mtime)
                    summary["updated"] += 1
            except Exception as exc:
                summary["errors"].append(f"{path_obj.name}: {exc}")

        for path_str, existing in tracked_by_path.items():
            if path_str not in disk_files:
                self._store.remove_training_file(existing.id)
                summary["removed"] += 1

        return summary

    def _add_from_training_folder(self, path: Path, mtime: float) -> TrainingFile:
        content = extract_text_from_file(path, path.suffix.lower())[:MAX_CONTENT_LENGTH]
        size_bytes = path.stat().st_size
        return self._store.add_training_file(
            filename=path.name,
            file_type=path.suffix.lstrip(".").lower(),
            size_bytes=size_bytes,
            content_text=content,
            source_path=str(path.resolve()),
            source_mtime=mtime,
        )

    def search(self, query: str, top_k: int = 5) -> List[TrainingFile]:
        return [doc for _score, doc in self.search_with_scores(query, top_k=top_k)]

    def search_with_scores(
        self, query: str, top_k: int = 5, embed_fn: Optional[Callable[[str], Optional[list]]] = None
    ) -> List[tuple]:
        """
        Busca documentos relevantes. Si se pasa `embed_fn` (típicamente
        el método `.embed()` del proveedor de IA activo) y devuelve un
        embedding real, se usa similitud semántica — entiende
        parafraseos y sinónimos que la búsqueda por keywords se pierde
        (ej. "restablecer mi clave" encuentra un documento que dice
        "cambiar la contraseña" aunque no compartan ninguna palabra).

        Si `embed_fn` no está disponible (proveedor sin soporte, sin
        conexión, o la llamada falla) o no encuentra nada por encima del
        umbral, cae de nuevo a la búsqueda por keywords de siempre — la
        Base de Conocimiento nunca deja de funcionar por esto.
        """
        if embed_fn is not None:
            semantic_results = self._search_semantic(query, embed_fn, top_k)
            if semantic_results:
                return semantic_results

        keywords = extract_keywords(query)
        if not keywords:
            return []
        return self._store.search_training_files_scored(keywords, top_k=top_k)

    def _search_semantic(
        self, query: str, embed_fn: Callable[[str], Optional[list]], top_k: int
    ) -> List[tuple]:
        query_embedding = embed_fn(query)
        if not query_embedding:
            return []

        docs_with_embeddings = self._store.list_training_files_with_embeddings()
        if not docs_with_embeddings:
            return []

        scored: List[tuple] = []
        for doc, embedding_json in docs_with_embeddings:
            doc_embedding = self._get_or_compute_embedding(doc, embedding_json, embed_fn)
            if not doc_embedding:
                continue
            similarity = cosine_similarity(query_embedding, doc_embedding)
            if similarity >= SEMANTIC_MATCH_THRESHOLD:
                scored.append((round(similarity, 4), doc))

        scored.sort(key=lambda pair: pair[0], reverse=True)
        return scored[:top_k]

    def _get_or_compute_embedding(
        self, doc: TrainingFile, embedding_json: str, embed_fn: Callable[[str], Optional[list]]
    ) -> Optional[list]:
        if embedding_json:
            try:
                return json.loads(embedding_json)
            except json.JSONDecodeError:
                pass  # embedding guardado corrupto: se recalcula abajo

        content = self._store.get_training_file_content(doc.id)
        if not content:
            return None

        embedding = embed_fn(content[:8000])
        if embedding:
            self._store.update_training_file_embedding(doc.id, json.dumps(embedding))
        return embedding

    def detect_ambiguous_matches(self, query: str, top_k: int = 5) -> List[TrainingFile]:
        scored = self.search_with_scores(query, top_k=top_k)
        if len(scored) < 2:
            return []

        top_score = scored[0][0]
        threshold = top_score * 0.85
        tied = [doc for score, doc in scored if score >= threshold]

        return tied if len(tied) >= 2 else []

    def remove_document(self, document_id: int) -> None:
        self._store.remove_training_file(document_id)

    def update_document(self, document_id: int, filename: Optional[str] = None) -> None:
        self._store.update_training_file(document_id, filename=filename)

    def build_context_snippet(self, matches: List[TrainingFile], max_chars_per_doc: int = 800) -> str:
        if not matches:
            return ""

        blocks = []
        for doc in matches:
            full_content = self._store.get_training_file_content(doc.id) or doc.content_preview
            blocks.append(f"--- {doc.filename} ---\n{full_content[:max_chars_per_doc]}")

        return "\n\n".join(blocks)
