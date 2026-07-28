from datetime import datetime
from pathlib import Path
from typing import List

from models.message import Message


class ExportError(Exception):
    pass


def _speaker_label(message: Message) -> str:
    return "Tú" if message.is_user else "Vicky"


def export_conversation_to_docx(title: str, messages: List[Message], output_path: Path) -> Path:
    try:
        import docx
        from docx.shared import Pt, RGBColor
    except ImportError as exc:
        raise ExportError(
            "Falta el paquete 'python-docx' para exportar a Word (pip install python-docx)."
        ) from exc

    document = docx.Document()
    document.add_heading(title or "Conversación", level=1)

    meta = document.add_paragraph(f"Exportado el {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    meta.runs[0].italic = True
    meta.runs[0].font.size = Pt(9)
    meta.runs[0].font.color.rgb = RGBColor(0x6E, 0x6E, 0x6E)

    for message in messages:
        header = document.add_paragraph()
        header_run = header.add_run(f"{_speaker_label(message)} · {message.timestamp}")
        header_run.bold = True
        header_run.font.size = Pt(10)
        if message.is_user:
            header_run.font.color.rgb = RGBColor(0xD8, 0x1F, 0x27)
        else:
            header_run.font.color.rgb = RGBColor(0x2B, 0x2B, 0x2B)

        for line in message.content.split("\n"):
            document.add_paragraph(line)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return output_path


def export_conversation_to_pdf(title: str, messages: List[Message], output_path: Path) -> Path:
    try:
        from reportlab.lib.colors import HexColor
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except ImportError as exc:
        raise ExportError(
            "Falta el paquete 'reportlab' para exportar a PDF (pip install reportlab)."
        ) from exc

    def escape(text: str) -> str:
        return (
            text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace("\n", "<br/>")
        )

    styles = getSampleStyleSheet()
    header_style = ParagraphStyle(
        "Header", parent=styles["Normal"], fontSize=10, spaceAfter=2, fontName="Helvetica-Bold"
    )
    meta_style = ParagraphStyle(
        "Meta", parent=styles["Normal"], fontSize=9, textColor=HexColor("#6E6E6E"), spaceAfter=14
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(str(output_path), pagesize=LETTER)

    story = [
        Paragraph(escape(title or "Conversación"), styles["Title"]),
        Paragraph(f"Exportado el {datetime.now().strftime('%d/%m/%Y %H:%M')}", meta_style),
    ]

    for message in messages:
        color = "#D81F27" if message.is_user else "#2B2B2B"
        header_style.textColor = HexColor(color)
        story.append(Paragraph(f"{_speaker_label(message)} · {message.timestamp}", header_style))
        story.append(Paragraph(escape(message.content), styles["BodyText"]))
        story.append(Spacer(1, 10))

    doc.build(story)
    return output_path
